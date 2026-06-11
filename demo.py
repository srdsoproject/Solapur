import streamlit as st
import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
import folium
from streamlit_folium import st_folium
import io

# Extra professional formatting imports
from docx import Document
from docx.shared import Inches

# ====================== PAGE CONFIG ======================
st.set_page_config(
    page_title="Solapur Division Equipment Management System",
    page_icon="🚉",
    layout="wide"
)

# ====================== STYLING ENGINE ======================
st.markdown("""
<style>
@keyframes fadeInUp { from { opacity: 0; transform: translateY(12px); } to { opacity: 1; transform: translateY(0); } }
@keyframes professionalFlash {
    0%, 100% { box-shadow: 0 4px 12px rgba(245, 158, 11, 0.1); border-top-color: #f59e0b; background: #fffbeb; }
    50% { box-shadow: 0 0 18px 6px rgba(245, 158, 11, 0.45); border-top-color: #d97706; background: #fef3c7; }
}
@keyframes glowGreen {
    0%, 100% { box-shadow: 0 4px 12px rgba(37, 99, 235, 0.08), 0 0 0 0 rgba(34, 197, 94, 0); }
    50% { box-shadow: 0 6px 20px rgba(34, 197, 94, 0.35), 0 0 14px 4px rgba(34, 197, 94, 0.2); }
}
@keyframes glowRed {
    0%, 100% { box-shadow: 0 4px 12px rgba(37, 99, 235, 0.08), 0 0 0 0 rgba(239, 68, 68, 0); }
    50% { box-shadow: 0 6px 20px rgba(239, 68, 68, 0.35), 0 0 14px 4px rgba(239, 68, 68, 0.2); }
}
.stApp { background-color: #f8fafc; color: #0f172a; }
html, body, [class*="css"] { font-family: "Segoe UI", -apple-system, BlinkMacSystemFont, sans-serif; }
.block-container { padding-top: 1.5rem; max-width: 1550px; animation: fadeInUp 0.4s cubic-bezier(0.16, 1, 0.3, 1) both; }
[data-testid="stSidebar"] { background: linear-gradient(180deg, #f0fdf4 0%, #dcfce7 100%); box-shadow: 4px 0 15px rgba(0,0,0,0.04); border-right: 1px solid #bbf7d0; }
[data-testid="stSidebar"] h3 { color: #14532d !important; }
[data-testid="stSidebar"] caption, [data-testid="stSidebar"] p, [data-testid="stSidebar"] span { color: #166534 !important; }
.header-box { background: linear-gradient(135deg, #1e3a8a 0%, #0369a1 100%); padding: 30px 40px; border-radius: 12px; margin-bottom: 25px; box-shadow: 0 10px 30px rgba(30, 58, 138, 0.12); border-left: 6px solid #38bdf8; }
.header-box h1 { color: #ffffff !important; font-size: 34px; font-weight: 800; margin: 0 0 6px 0; letter-spacing: -0.5px; }
.header-box h4 { color: #bae6fd !important; font-weight: 400; margin: 0; letter-spacing: 0.2px; }
.rail-station-wrapper { background: #ffffff; border: 1px solid #e2e8f0; border-radius: 12px; padding: 22px; margin-bottom: 25px; box-shadow: 0 4px 14px rgba(0,0,0,0.02); }
.station-title-strip { font-size: 20px; font-weight: 700; color: #1e3a8a !important; border-bottom: 2px solid #f1f5f9; padding-bottom: 10px; margin-bottom: 18px; }
.metric-box { background: #ffffff; border: 1px solid #e2e8f0; border-radius: 8px; padding: 16px 14px; text-align: center; transition: all 0.2s ease; height: 100%; display: flex; flex-direction: column; justify-content: center; align-items: center; gap: 6px; }
.metric-box.active-assets { border-top: 4px solid #2563eb; background: #f8fafc; }
.metric-box.active-assets:hover { transform: scale(1.03); background: #eff6ff; }
.metric-box.infra-asset { border-top: 4px solid #0284c7; background: #f0f9ff; }
.metric-box.infra-asset:hover { transform: scale(1.03); background: #e0f2fe; }
.metric-box.flag-green-glow { border-top: 4px solid #22c55e !important; background: #f0fdf4 !important; animation: glowGreen 2.5s infinite ease-in-out; }
.metric-box.flag-red-glow { border-top: 4px solid #ef4444 !important; background: #fef2f2 !important; animation: glowRed 2.5s infinite ease-in-out; }
.metric-box.flag-flash-glow { border-top: 4px solid #f59e0b !important; animation: professionalFlash 2s infinite ease-in-out; }
.metric-box.zero-assets { border-top: 4px solid #94a3b8; background: #ffffff; opacity: 0.75; }
.metric-label { color: #334155 !important; font-size: 14px !important; font-weight: 600 !important; text-transform: capitalize; text-align: center; line-height: 1.3; margin: 0; }
.metric-value { color: #1e3a8a !important; font-size: 17px !important; font-weight: 700 !important; margin: 0; }
.metric-value.zero { color: #94a3b8 !important; font-weight: 600; }
.export-panel { background: #f1f5f9; padding: 15px; border-radius: 8px; border: 1px solid #cbd5e1; margin-top: 10px; margin-bottom: 20px;}

/* MOBILE-RESPONSIVE VIEWPORT ENGINE (POINT 5) */
@media (max-width: 768px) {
    .block-container {
        padding-left: 0.5rem !important;
        padding-right: 0.5rem !important;
        padding-top: 0.5rem !important;
    }
    .header-box {
        padding: 18px 20px !important;
        margin-bottom: 15px !important;
    }
    .header-box h1 {
        font-size: 22px !important;
        line-height: 1.2;
    }
    .header-box h4 {
        font-size: 13px !important;
    }
    .rail-station-wrapper {
        padding: 12px !important;
        margin-bottom: 15px !important;
    }
    .station-title-strip {
        font-size: 16px !important;
        margin-bottom: 12px !important;
    }
    [data-testid="stHorizontalBlock"] {
        gap: 8px !important;
    }
}
</style>
""", unsafe_allow_html=True)

# ====================== LOGIN SERVICE ======================
def login():
    if st.session_state.get("authenticated"):
        return True
    
    _, col2, _ = st.columns([1, 1.6, 1])
    with col2:
        st.write("")
        st.markdown("""
        <div style="background: white; padding: 35px; border-radius: 12px; box-shadow: 0 10px 25px rgba(0,0,0,0.04); border: 1px solid #e2e8f0; text-align: center;">
            <h2 style="margin: 0 0 5px 0; color: #1e3a8a; font-weight:800;">🚉 Solapur Division</h2>
            <p style="color: #64748b; margin: 0; font-size:14px; font-weight:500;">Safety & Engineering Branch Asset Tracking Portal</p>
        </div>
        """, unsafe_allow_html=True)
        
        with st.form("login_form"):
            username = st.text_input("Operator User ID")
            password = st.text_input("Security Access Password", type="password")
            if st.form_submit_button("Secure Log In", use_container_width=True):
                try:
                    if username == st.secrets["APP_USERNAME"] and password == st.secrets["APP_PASSWORD"]:
                        st.session_state.authenticated = True
                        st.rerun()
                    else:
                        st.error("❌ Invalid User ID or Password credential.")
                except KeyError:
                    st.error("🔑 Context configuration parameters are missing.")
    return False

if not login():
    st.stop()

# ====================== RESILIENT DATA PIPELINE ENGINE (POINT 5) ======================
@st.cache_data(ttl=600, show_spinner="Fetching Data Matrix...")
def fetch_from_google(sheet_id_key, sheet_name_key):
    """Direct network connection layer to Cloud Database."""
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"],
        scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    )
    df = pd.DataFrame(gspread.authorize(creds).open_by_key(st.secrets[sheet_id_key]).worksheet(st.secrets[sheet_name_key]).get_all_records())
    df.columns = df.columns.str.strip()
    return df

def load_secure_sheet(sheet_id_key, sheet_name_key):
    """Smart cache-first manager with an un-crashable network fallback pipeline."""
    backup_key = f"backup_{sheet_id_key}_{sheet_name_key}"
    
    try:
        # Request fresh data (or pull straight from Streamlit cache memory if inside TTL window)
        df = fetch_from_google(sheet_id_key, sheet_name_key)
        
        # Commit a copy to session memory state as the ultimate network loss fallback
        st.session_state[backup_key] = df
        return df
        
    except Exception:
        # Network signal is lost! Check if an offline session backup exists
        if backup_key in st.session_state:
            return st.session_state[backup_key]
        else:
            # Absolute worst-case scenario: No internet and no prior cached memory
            st.error("🚨 Connection failure. No offline cached data found yet for this department.")
            return pd.DataFrame()

# ====================== AUTOMATED DYNAMIC GRID COMPONENT ======================
def render_dynamic_grid(df, icon_emoji="📦", col_layout=4, is_inventory=False):
    if df.empty:
        return
        
    all_headers = list(df.columns)
    title_column = all_headers[0]  
    detail_cols = [c for c in all_headers if c not in [title_column, "LATITUDE", "LONGITUDE"]]
    
    for _, row in df.iterrows():
        title_value = row.get(title_column, "Unknown Node")
        
        st.markdown(f"""
        <div class="rail-station-wrapper">
            <div class="station-title-strip">{icon_emoji} {title_column}: <b>{title_value}</b></div>
        """, unsafe_allow_html=True)
        
        cols = st.columns(col_layout)
        for idx, col in enumerate(detail_cols):
            raw = row.get(col, 0 if is_inventory else "N/A")
            val = "N/A" if (pd.isna(raw) or raw == "") else raw
            
            if is_inventory:
                try: val = int(float(raw)) if raw != "" and not pd.isna(raw) else 0
                except: val = raw
                box_style = "zero-assets" if val == 0 else "active-assets"
                val_style = "zero" if val == 0 else "active"
                
                if val != 0:
                    c_clean = str(col).lower()
                    if "green" in c_clean and "flag" in c_clean: box_style += " flag-green-glow"
                    elif "red" in c_clean and "flag" in c_clean: box_style += " flag-red-glow"
                    elif "led" in c_clean: box_style += " flag-flash-glow"
            else:
                box_style = "infra-asset"
                val_style = "active"
                
            with cols[idx % col_layout]:
                st.markdown(f"""
                <div class="metric-box {box_style}">
                    <div class="metric-label">{col}</div>
                    <div class="metric-value {val_style}">{val}</div>
                </div>
                """, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

# ====================== PAGINATION HELPER FUNCTION ======================
def handle_pagination(key_prefix, total_items, items_per_page):
    page_key = f"{key_prefix}_page"
    if page_key not in st.session_state:
        st.session_state[page_key] = 0

    total_pages = max(1, (total_items + items_per_page - 1) // items_per_page)
    
    if st.session_state[page_key] >= total_pages:
        st.session_state[page_key] = total_pages - 1

    col_p1, col_p2, col_p3 = st.columns([1, 2, 1])
    
    with col_p1:
        if st.button("⬅️ Previous", key=f"{key_prefix}_prev", disabled=(st.session_state[page_key] == 0), use_container_width=True):
            st.session_state[page_key] -= 1
            st.rerun()
            
    with col_p2:
        st.markdown(f"<p style='text-align: center; font-weight: 600; padding-top:6px; color: #1e3a8a;'>Page {st.session_state[page_key] + 1} of {total_pages} (Total Records: {total_items})</p>", unsafe_allow_html=True)
        
    with col_p3:
        if st.button("Next ➡️", key=f"{key_prefix}_next", disabled=(st.session_state[page_key] >= total_pages - 1), use_container_width=True):
            st.session_state[page_key] += 1
            st.rerun()

    start_idx = st.session_state[page_key] * items_per_page
    end_idx = start_idx + items_per_page
    return start_idx, end_idx

# ====================== ADVANCED MULTI-FORMAT EXPORT ENGINE ======================
def render_export_desk(df, prefix_name):
    """Renders an intuitive, robust data exportation dashboard module component."""
    if df.empty:
        return

    st.markdown("""<div class="export-panel">💾 <b>System Document & Export Desk</b></div>""", unsafe_allow_html=True)
    base_filename = f"Solapur_Division_{prefix_name.replace(' ', '_')}_Report"
    
    col_f1, col_f2 = st.columns([1, 2])
    with col_f1:
        format_choice = st.selectbox(
            "Select Target Output Format:",
            ["Excel Worksheets (.xlsx)", "Word Document (.docx)", "CSV Matrix (.csv)", "Printable Report Summary (PDF / HTML)"],
            key=f"format_{prefix_name}"
        )
        
    with col_f2:
        st.write("") 
        st.write("") 
        
        # 1. EXCEL DOCUMENT ENGINE
        if "Excel" in format_choice:
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name="Asset Inventory Log")
            st.download_button(
                label="📥 Save File as Complete Excel Document",
                data=buffer.getvalue(),
                file_name=f"{base_filename}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
            
        # 2. WORD NARRATIVE ENGINE
        elif "Word" in format_choice:
            doc = Document()
            doc.add_heading(f"Solapur Division Asset Inventory Report - {prefix_name}", level=1)
            doc.add_paragraph("Generated dynamically via Safety & Engineering Branch Management Portal.")
            
            t = doc.add_table(rows=1, cols=len(df.columns))
            t.style = 'Light Shading Accent 1'
            hdr_cells = t.rows[0].cells
            for i, col_name in enumerate(df.columns):
                hdr_cells[i].text = str(col_name)
                
            for _, row in df.head(100).iterrows(): 
                row_cells = t.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
                    
            buffer = io.BytesIO()
            doc.save(buffer)
            st.download_button(
                label="📥 Save File as Microsoft Word Narrative Report",
                data=buffer.getvalue(),
                file_name=f"{base_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        # 3. CSV ENGINE
        elif "CSV" in format_choice:
            csv_bytes = df.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📥 Save File as Plain Text CSV Matrix",
                data=csv_bytes,
                file_name=f"{base_filename}.csv",
                mime="text/csv",
                use_container_width=True
            )
            
        # 4. PRINT REPORT OVERVIEW (FOR PDF GENERATION VIA CTRL+P)
        elif "Printable" in format_choice:
            html_table = df.to_html(index=False, classes='table table-striped')
            html_blob = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 30px; }}
                    h2 {{ color: #1e3a8a; }}
                    table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
                    th {{ background-color: #0369a1; color: white; padding: 10px; text-align: left; }}
                    td {{ padding: 8px; border: 1px solid #cbd5e1; }}
                    tr:nth-child(even) {{ background-color: #f8fafc; }}
                </style>
            </head>
            <body>
                <h2>🚉 Solapur Division - {prefix_name} Structural Logs</h2>
                <p>Official Safety Desk Output Log. Press <b>Ctrl + P</b> to save directly as an official corporate PDF format profile.</p>
                {html_table}
            </body>
            </html>
            """
            st.download_button(
                label="📥 Save File as Web Report / PDF (Open HTML & Press Ctrl+P)",
                data=html_blob.encode('utf-8'),
                file_name=f"{base_filename}.html",
                mime="text/html",
                use_container_width=True
            )
    st.markdown("---")

# ====================== APPLICATION MAIN CORE ======================
def main_portal():
    st.markdown("""
    <div class="header-box">
        <h1>🚉 Solapur Division Inventory Management Dashboard</h1>
        <h4>An initiative by Safety Branch Solapur Division</h4>
    </div>
    """, unsafe_allow_html=True)

    # --- Sidebar Console Operations ---
    st.sidebar.markdown("### 🖥️ Controller Desk")
    st.sidebar.caption("Google Sheet status: Connected")
    
    if st.sidebar.button("🔄 Sync Google Sheets & Clear Cache", use_container_width=True):
        fetch_from_google.clear()
        for key in list(st.session_state.keys()):
            if "_page" in key:
                st.session_state[key] = 0
        st.rerun()
        
    if st.sidebar.button("🚪 Terminate Session", use_container_width=True):
        st.session_state.clear()
        st.rerun()

    # --- STATE PRESERVING SIDEBAR ROUTER ---
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🗺️ Department Navigation")
    
    available_departments = ["🚞 Operating", "🛠️ Engineering", "👨‍✈️ TRO", "⚡ TRD", "🚧 LC Gates"]
    
    if "selected_dept" not in st.session_state:
        st.session_state.selected_dept = "🚞 Operating"
        
    chosen_dept = st.sidebar.radio(
        "Choose Department View:",
        available_departments,
        index=available_departments.index(st.session_state.selected_dept)
    )
    st.session_state.selected_dept = chosen_dept
    
    max_cards = 25

    # ---------------------------------------------------------------
    # BRANCH VIEW 1: OPERATING
    # ---------------------------------------------------------------
    if st.session_state.selected_dept == "🚞 Operating":
        st.markdown("## 🚞 Operating Assets Matrix")
        df_eq = load_secure_sheet("SHEET_ID", "SHEET_NAME")
        if df_eq.empty:
            st.warning("⚠️ No active rows found inside Equipment Inventory database.")
        else:
            first_col = df_eq.columns[0]
            other_cols = [c for c in df_eq.columns if c not in [first_col, "LATITUDE", "LONGITUDE"]]
            
            kpi1, kpi2, kpi3 = st.columns(3)
            num_df = df_eq[other_cols].apply(pd.to_numeric, errors='coerce').fillna(0)
            kpi1.metric("🚏 Active Nodes", f"{len(df_eq)} Locations")
            kpi2.metric("🛠️ Tracked Column Heads", f"{len(df_eq.columns)} Metrics")
            kpi3.metric("📦 Gross Physical Units", f"{int(num_df.sum().sum()):,} Elements")
            
            if "LATITUDE" in df_eq.columns and "LONGITUDE" in df_eq.columns:
                with st.expander("🗺️ Division GIS Spatial Mapping Overview", expanded=False):
                    map_lat = pd.to_numeric(df_eq["LATITUDE"], errors='coerce').dropna().mean()
                    map_lon = pd.to_numeric(df_eq["LONGITUDE"], errors='coerce').dropna().mean()
                    if not (pd.isna(map_lat) or pd.isna(map_lon)):
                        m = folium.Map(location=[map_lat, map_lon], zoom_start=8)
                        for _, r in df_eq.head(40).iterrows():
                            try:
                                lat, lon = float(r["LATITUDE"]), float(r["LONGITUDE"])
                                if not (pd.isna(lat) or pd.isna(lon)):
                                    folium.Marker([lat, lon], tooltip=f"Location: {r.get(first_col, 'Unknown')}", icon=folium.Icon(color="blue", icon="train", prefix="fa")).add_to(m)
                            except: continue
                        st_folium(m, width="100%", height=320, key="global_gis_map", returned_objects=[])

            search_eq = st.text_input("🔍 Operational Search Desk Filter", placeholder=f"Search by {first_col}...", key="search_eq")
            fil_df = df_eq[df_eq[first_col].astype(str).str.contains(search_eq, case=False, na=False)] if search_eq else df_eq
            
            render_export_desk(fil_df, "Operating")
            start, end = handle_pagination("operating", len(fil_df), max_cards)
            render_dynamic_grid(fil_df.iloc[start:end], icon_emoji="🚞", is_inventory=True)

    # ---------------------------------------------------------------
    # BRANCH VIEW 2: ENGINEERING
    # ---------------------------------------------------------------
    elif st.session_state.selected_dept == "🛠️ Engineering":
        st.markdown("## 🛠️ Engineering Infrastructure Matrix")
        df_tr = load_secure_sheet("NEW_SHEET_ID", "NEW_SHEET_NAME")
        if df_tr.empty:
            st.warning("⚠️ No active infrastructure data elements found.")
        else:
            first_col = df_tr.columns[0]
            kpi_t1, kpi_t2 = st.columns(2)
            kpi_t1.metric("🛤️ Total Matrix Items", f"{len(df_tr)} Records")
            kpi_t2.metric("📋 Structural Column Metrics", f"{len(df_tr.columns)} Heads")
            
            search_tr = st.text_input("🔍 Engineering Search Desk Filter", placeholder=f"Search by {first_col}...", key="search_tr")
            fil_df_tr = df_tr.copy()
            if search_tr:
                fil_df_tr = fil_df_tr[fil_df_tr[first_col].astype(str).str.contains(search_tr, case=False, na=False)]
                
            render_export_desk(fil_df_tr, "Engineering")
            start, end = handle_pagination("engineering", len(fil_df_tr), max_cards)
            render_dynamic_grid(fil_df_tr.iloc[start:end], icon_emoji="🛠️")

    # ---------------------------------------------------------------
    # BRANCH VIEW 3: TRO
    # ---------------------------------------------------------------
    elif st.session_state.selected_dept == "👨‍✈️ TRO":
        st.markdown("## 👨‍✈️ TRO Facility Operations Matrix")
        df_tro = load_secure_sheet("TRO", "NEW_SHEET_NAME_TRO")
        if df_tro.empty:
            st.warning("⚠️ No active crew facility elements found inside the TRO matrix.")
        else:
            first_col = df_tro.columns[0]
            kpi_tro1, kpi_tro2 = st.columns(2)
            kpi_tro1.metric("🏢 Facilities Logged", f"{len(df_tro)} Units")
            kpi_tro2.metric("📋 Automated System Headers", f"{len(df_tro.columns)} Fields")
            
            search_tro = st.text_input("🔍 Crew Management Search Filter", placeholder=f"Search by {first_col}...", key="search_tro")
            fil_df_tro = df_tro.copy()
            if search_tro:
                fil_df_tro = fil_df_tro[fil_df_tro[first_col].astype(str).str.contains(search_tro, case=False, na=False)]
                
            render_export_desk(fil_df_tro, "TRO")
            start, end = handle_pagination("tro", len(fil_df_tro), max_cards)
            render_dynamic_grid(fil_df_tro.iloc[start:end], icon_emoji="👨‍✈️", col_layout=2)

    # ---------------------------------------------------------------
    # BRANCH VIEW 4: TRD
    # ---------------------------------------------------------------
    elif st.session_state.selected_dept == "⚡ TRD":
        st.markdown("## ⚡ TRD Power Allocation Distribution Matrix")
        df_trd = load_secure_sheet("SHEET_ID_TRD", "SHEET_NAME_TRD")
        if df_trd.empty:
            st.warning("⚠️ No active Traction Distribution (TRD) data elements found.")
        else:
            first_col = df_trd.columns[0]
            kpi_trd1, kpi_trd2 = st.columns(2)
            kpi_trd1.metric("⚡ TRD Operations Sectors", f"{len(df_trd)} Monitored Records")
            kpi_trd2.metric("📋 Tracked System Headers", f"{len(df_trd.columns)} Fields")
            
            search_trd = st.text_input("🔍 TRD Technical Asset Search Filter", placeholder=f"Search by {first_col}...", key="search_trd")
            fil_df_trd = df_trd.copy()
            if search_trd:
                fil_df_trd = fil_df_trd[fil_df_trd[first_col].astype(str).str.contains(search_trd, case=False, na=False)]
                
            render_export_desk(fil_df_trd, "TRD")
            start, end = handle_pagination("trd", len(fil_df_trd), max_cards)
            render_dynamic_grid(fil_df_trd.iloc[start:end], icon_emoji="⚡")

    # ---------------------------------------------------------------
    # BRANCH VIEW 5: LC GATES
    # ---------------------------------------------------------------
    elif st.session_state.selected_dept == "🚧 LC Gates":
        st.markdown("## 🚧 Level Crossing Gates Registry Matrix")
        df_lc = load_secure_sheet("SHEET_ID_LC", "SHEET_NAME_LC")
        if df_lc.empty:
            st.warning("⚠️ No active Level Crossing (LC) Gate structural elements found.")
        else:
            first_col = df_lc.columns[0]
            kpi_lc1, kpi_lc2 = st.columns(2)
            kpi_lc1.metric("🚧 Level Crossing Records", f"{len(df_lc)} Active Crossings")
            kpi_lc2.metric("📋 Automated Gate Fields", f"{len(df_lc.columns)} Column heads")
            
            search_lc = st.text_input("🔍 Crossing Point Search Filter", placeholder=f"Search by {first_col}...", key="search_lc")
            fil_df_lc = df_lc.copy()
            if search_lc:
                fil_df_lc = fil_df_lc[fil_df_lc[first_col].astype(str).str.contains(search_lc, case=False, na=False)]
                
            render_export_desk(fil_df_lc, "LC_Gates")
            start, end = handle_pagination("lc_gates", len(fil_df_lc), max_cards)
            render_dynamic_grid(fil_df_lc.iloc[start:end], icon_emoji="🚧")

if __name__ == "__main__":
    main_portal()
